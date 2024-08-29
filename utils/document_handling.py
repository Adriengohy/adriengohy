import docx
from datetime import datetime


def _add_header(doc, header, level):
    doc.add_heading(header, level=level)
    return doc


def _add_paragraph(doc, text):
    doc.add_paragraph(text)
    return doc


def _add_page_break(doc):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_break(docx.enum.text.WD_BREAK.PAGE)
    
    return doc


def _get_today_date_formatted():
    today = datetime.now()
    formatted_date = today.strftime("%d %B %Y")
    return formatted_date


def get_input_file(input_file):
    doc = docx.Document(input_file)
    return doc


def fill_in_front_page(doc, title):
    for paragraph in doc.paragraphs:
        if "[title]" in paragraph.text:
            for run in paragraph.runs:
                run.text = run.text.replace("[title]", title)
        
        if "[DD Month YYYY]" in paragraph.text:
            for run in paragraph.runs:
                run.text = run.text.replace("[DD Month YYYY]", _get_today_date_formatted())
    return doc


def write_document(doc, section, level=1): # recursieve functie die doorheen de tekst json gaat en deze in de doc invult
    doc = _add_header(doc, section["title"], level)
    for text in section["content"]:
        doc = _add_paragraph(doc, text)
    for subsection in section["sections"]:
        doc = write_document(doc, subsection, level+1)
    
    if level == 1:
        doc = _add_page_break(doc)
    return doc
